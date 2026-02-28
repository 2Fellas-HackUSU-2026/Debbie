import os
import uuid
import json
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
from dotenv import load_dotenv

load_dotenv(os.path.join(os.path.dirname(__file__), "..", "config", ".env"))

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

UPLOAD_DIR = "api/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class Selection(BaseModel):
    id: str
    variable_name: str
    purpose: Optional[str] = None
    is_checkbox: bool = False

class ProcessRequest(BaseModel):
    doc_id: str
    selections: List[Selection]

def get_alignment_string(alignment):
    if alignment == WD_ALIGN_PARAGRAPH.CENTER: return "center"
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT: return "right"
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY: return "justify"
    return "left"

def parse_docx(filepath: str) -> Dict[str, Any]:
    doc = Document(filepath)
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        runs = [{"text": r.text, "bold": r.bold, "italic": r.italic, "underline": r.underline} for r in p.runs]
        paragraphs.append({
            "id": f"p_{i}",
            "text": p.text,
            "style": p.style.name,
            "alignment": get_alignment_string(p.alignment),
            "runs": runs,
            "has_checkbox": "☐" in p.text or "☑" in p.text
        })

    tables = []
    for i, t in enumerate(doc.tables):
        rows = []
        for r_idx, row in enumerate(t.rows):
            cells = []
            for c_idx, cell in enumerate(row.cells):
                cell_paras = []
                for p_idx, p in enumerate(cell.paragraphs):
                    runs = [{"text": r.text, "bold": r.bold, "italic": r.italic, "underline": r.underline} for r in p.runs]
                    cell_paras.append({
                        "id": f"t_{i}_r_{r_idx}_c_{c_idx}_p_{p_idx}",
                        "text": p.text,
                        "alignment": get_alignment_string(p.alignment),
                        "runs": runs,
                        "has_checkbox": "☐" in p.text or "☑" in p.text
                    })
                cells.append({"id": f"t_{i}_r_{r_idx}_c_{c_idx}", "paragraphs": cell_paras})
            rows.append(cells)
        tables.append({"id": f"t_{i}", "rows": rows, "num_columns": len(t.columns)})

    return {"paragraphs": paragraphs, "tables": tables}

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"): raise HTTPException(status_code=400, detail="Only .docx files are supported")
    doc_id = str(uuid.uuid4())
    filepath = os.path.join(UPLOAD_DIR, f"{doc_id}.docx")
    with open(filepath, "wb") as b: b.write(await file.read())
    structure = parse_docx(filepath)
    with open(os.path.join(UPLOAD_DIR, f"{doc_id}.json"), "w") as f: json.dump(structure, f)
    return {"doc_id": doc_id, "structure": structure}

@app.get("/suggest/{doc_id}")
async def suggest_regions(doc_id: str):
    struct_path = os.path.join(UPLOAD_DIR, f"{doc_id}.json")
    if not os.path.exists(struct_path): raise HTTPException(status_code=404, detail="Not found")
    with open(struct_path, "r") as f: structure = json.load(f)
    prompt = "Identify fillable fields or checkboxes. Return JSON key 'suggestions' as list of {id, suggested_name}.\n\n" + json.dumps(structure)[:8000]
    api_key = os.getenv("OPEN_AI_KEY") or os.getenv("OPENAI_API_KEY")
    client = openai.OpenAI(api_key=api_key)
    try:
        resp = client.chat.completions.create(model="gpt-4o", messages=[{"role": "system", "content": "Assistant"}, {"role": "user", "content": prompt}], response_format={"type": "json_object"})
        return json.loads(resp.choices[0].message.content)
    except: return {"suggestions": []}

@app.post("/process")
async def process_document(request: ProcessRequest):
    doc_path = os.path.join(UPLOAD_DIR, f"{request.doc_id}.docx")
    if not os.path.exists(doc_path): raise HTTPException(status_code=404, detail="Not found")

    doc = Document(doc_path)
    mapping = {}

    # Track which cells/paragraphs are overridden by individual selections
    overridden = set()

    # Sort selections: columns first, then individual cells (to allow overrides)
    sorted_selections = sorted(request.selections, key=lambda s: 0 if "col" in s.id else 1)

    for selection in sorted_selections:
        mapping[selection.variable_name] = {"id": selection.id, "purpose": selection.purpose}
        placeholder = f"{{{{ {selection.variable_name} }}}}"

        if selection.id.startswith("p_"):
            p_idx = int(selection.id.split("_")[1])
            doc.paragraphs[p_idx].text = placeholder
        elif selection.id.startswith("t_"):
            parts = selection.id.split("_")
            t_idx = int(parts[1])
            table = doc.tables[t_idx]
            if "col" in selection.id:
                col_idx = int(parts[3])
                for r_idx, row in enumerate(table.rows):
                    cell_id = f"t_{t_idx}_r_{r_idx}_c_{col_idx}"
                    if cell_id not in overridden:
                        row.cells[col_idx].text = f"{{{{ {selection.variable_name}_{r_idx} }}}}"
            else:
                r_idx, c_idx = int(parts[3]), int(parts[5])
                overridden.add(selection.id)
                table.rows[r_idx].cells[c_idx].text = placeholder

    output_path = os.path.join(UPLOAD_DIR, f"{request.doc_id}_templated.docx")
    doc.save(output_path)
    with open(os.path.join(UPLOAD_DIR, f"{request.doc_id}_mapping.json"), "w") as f: json.dump(mapping, f)
    return {"message": "Success", "download_url": f"/download/{request.doc_id}"}

@app.get("/download/{doc_id}")
async def download_document(doc_id: str):
    path = os.path.join(UPLOAD_DIR, f"{doc_id}_templated.docx")
    return FileResponse(path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="templated.docx")

@app.get("/")
async def read_index(): return FileResponse("web/index.html")
if os.path.exists("web"): app.mount("/static", StaticFiles(directory="web"), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
