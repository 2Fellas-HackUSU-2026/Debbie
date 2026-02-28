import os
import uuid
import json
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
import openai
from dotenv import load_dotenv

load_dotenv(os.path.join(os.path.dirname(__file__), "..", "config", ".env"))

app = FastAPI()

# Enable CORS for frontend interaction
# In production, replace ["*"] with specific origins
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

# In-memory storage for document structures (cleared on restart)
document_structures: Dict[str, Dict[str, Any]] = {}

UPLOAD_DIR = "api/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


class Selection(BaseModel):
    id: str  # e.g., "p_0", "t_0_r_1_c_2", or "t_0_col_2"
    variable_name: str
    description: Optional[str] = None


class ProcessRequest(BaseModel):
    doc_id: str
    selections: List[Selection]


def extract_run_formatting(run) -> Dict[str, bool]:
    return {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline
    }


def get_paragraph_alignment(p) -> str:
    alignment = "left"
    if p.alignment:
        if p.alignment == 1:
            alignment = "center"
        elif p.alignment == 2:
            alignment = "right"
        elif p.alignment == 3:
            alignment = "justify"
    return alignment


def process_text_runs(p) -> List[Dict[str, Any]]:
    runs = []
    for run in p.runs:
        runs.append({
            "text": run.text,
            "formatting": extract_run_formatting(run)
        })
    return runs


def extract_checkboxes(text: str) -> bool:
    checkboxes = ["☐", "☑", "☒"]
    for cb in checkboxes:
        if cb in text:
            return True
    return False


def parse_docx(filepath: str) -> Dict[str, Any]:
    doc = Document(filepath)
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        has_checkbox = extract_checkboxes(text)
        paragraphs.append({
            "id": f"p_{i}",
            "text": text,
            "style": p.style.name,
            "alignment": get_paragraph_alignment(p),
            "runs": process_text_runs(p),
            "has_checkbox": has_checkbox,
            "is_blank": not text.strip()
        })

    tables = []
    for i, t in enumerate(doc.tables):
        rows = []
        for r_idx, row in enumerate(t.rows):
            cells = []
            for c_idx, cell in enumerate(row.cells):
                text = cell.text
                has_checkbox = extract_checkboxes(text)

                # We need cell paragraphs formatting as well
                cell_paragraphs = []
                for cp in cell.paragraphs:
                    cell_paragraphs.append({
                        "text": cp.text,
                        "style": cp.style.name,
                        "alignment": get_paragraph_alignment(cp),
                        "runs": process_text_runs(cp),
                    })

                cells.append({
                    "id": f"t_{i}_r_{r_idx}_c_{c_idx}",
                    "text": text,
                    "paragraphs": cell_paragraphs,
                    "has_checkbox": has_checkbox,
                    "is_blank": not text.strip()
                })
            rows.append(cells)
        tables.append({
            "id": f"t_{i}",
            "rows": rows,
            "num_columns": len(t.columns)
        })

    return {"paragraphs": paragraphs, "tables": tables}


@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(
            status_code=400, detail="Only .docx files are supported")

    doc_id = str(uuid.uuid4())
    filepath = os.path.join(UPLOAD_DIR, f"{doc_id}.docx")

    with open(filepath, "wb") as buffer:
        buffer.write(await file.read())

    structure = parse_docx(filepath)

    # Store structure in memory
    document_structures[doc_id] = structure

    return {"doc_id": doc_id, "structure": structure}


@app.get("/suggest/{doc_id}")
async def suggest_regions(doc_id: str):
    if doc_id not in document_structures:
        raise HTTPException(status_code=404, detail="Document not found")

    structure = document_structures[doc_id]

    # Filter the structure to only include areas WITH text (non-blank) to suggest variable names
    # Note: the user asked to ONLY add suggestions to areas that HAVE TEXT on the document.
    non_blank_paragraphs = [p for p in structure.get(
        'paragraphs', []) if not p.get('is_blank', True)]
    non_blank_tables = []
    for t in structure.get('tables', []):
        non_blank_rows = []
        for row in t.get('rows', []):
            non_blank_cells = [c for c in row if not c.get('is_blank', True)]
            if non_blank_cells:
                non_blank_rows.append(non_blank_cells)
        if non_blank_rows:
            non_blank_tables.append({"id": t['id'], "rows": non_blank_rows})

    filtered_structure = {
        "paragraphs": non_blank_paragraphs, "tables": non_blank_tables}

    # Prepare prompt for OpenAI
    prompt_content = "Identify potential fillable fields in this Word document structure. ONLY CONSIDER THE PROVIDED TEXT (WHICH IS NON-BLANK). "
    prompt_content += "Suggest a clean, snake_case variable name based on the content of the field itself and its surrounding context (e.g. headers). "
    prompt_content += "Return a JSON object with a key 'suggestions' which is a list of objects with 'id' (the identifier from the structure) and 'suggested_name'.\n\n"
    prompt_content += json.dumps(filtered_structure)[:8000]

    api_key = os.getenv("OPEN_AI_KEY") or os.getenv("OPENAI_API_KEY")
    client = openai.OpenAI(api_key=api_key)

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a document analysis assistant. Return only valid JSON."},
                {"role": "user", "content": prompt_content}
            ],
            response_format={"type": "json_object"}
        )
        suggestions = json.loads(response.choices[0].message.content)
        return suggestions
    except Exception as e:
        return {"suggestions": [], "error": str(e)}


class SuggestionRequest(BaseModel):
    text: str


@app.post("/suggest")
async def suggest_variable_name(request: SuggestionRequest):
    """Generate a variable name suggestion based on the provided text."""
    text = request.text.strip()

    if not text:
        raise HTTPException(status_code=400, detail="Text cannot be empty")

    api_key = os.getenv("OPEN_AI_KEY") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=500, detail="OpenAI API key not configured")

    client = openai.OpenAI(api_key=api_key)

    try:
        prompt_content = f"""Based on the following text content from document fields (possibly multiple cells in a column), suggest a single clean, descriptive variable name in snake_case format that represents what ALL these fields are for.
        
Text content from fields:
{text}

Return a JSON object with a single key 'suggestion' containing only the variable name (e.g., "project_name", "employee_id", "date_submitted", "item_description").
The variable name should be concise, descriptive, and represent the overall purpose of these fields."""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a variable naming assistant. Return only valid JSON with a 'suggestion' key."},
                {"role": "user", "content": prompt_content}
            ],
            response_format={"type": "json_object"}
        )

        result = json.loads(response.choices[0].message.content)
        return result
    except Exception as e:
        return {"suggestion": "", "error": str(e)}


@app.post("/process")
async def process_document(request: ProcessRequest):
    doc_path = os.path.join(UPLOAD_DIR, f"{request.doc_id}.docx")
    if not os.path.exists(doc_path):
        raise HTTPException(status_code=404, detail="Document not found")

    doc = Document(doc_path)

    # We want a flat dictionary for data.json: { "variable_name": "description (and checkbox context)" }
    data_dict = {}

    # Get structure from memory
    if request.doc_id not in document_structures:
        raise HTTPException(status_code=404, detail="Document not found")
    
    structure = document_structures[request.doc_id]
    
    # Read structure to identify checkboxes
    checkbox_ids = set()
    for p in structure.get("paragraphs", []):
        if p.get("has_checkbox"):
            checkbox_ids.add(p["id"])
    for t in structure.get("tables", []):
        for row in t.get("rows", []):
            for cell in row:
                if cell.get("has_checkbox"):
                    checkbox_ids.add(cell["id"])

    # Collect individual cell assignments first, so they take precedence over columns
    # We'll map cell ID to its assigned variable
    cell_overrides = {}
    col_assignments = []

    for selection in request.selections:
        is_checkbox = False

        # Check if selection corresponds to a checkbox
        if selection.id in checkbox_ids:
            is_checkbox = True
        elif selection.id.startswith("t_") and "col" in selection.id:
            # Need to check if any cell in the column has a checkbox (complex, simplify to description)
            # Actually, col selections don't map directly to a single element's checkbox flag easily.
            pass

        context_val = selection.description or ""
        if is_checkbox:
            context_val = context_val + \
                " (Field is a checkbox)" if context_val else "Field is a checkbox"

        # Determine behavior based on selection type
        if selection.id.startswith("t_") and "col" in selection.id:
            col_assignments.append(selection)
        else:
            cell_overrides[selection.id] = selection.variable_name
            data_dict[selection.variable_name] = context_val

    # Also add column assignments to data_dict with all their individual variables
    for selection in col_assignments:
        context_val = selection.description or ""
        # For columns, we need to add all individual row variables
        # Parse the column info to get number of rows
        parts = selection.id.split("_")
        if len(parts) >= 4:
            t_idx = int(parts[1])
            col_idx = int(parts[3])

            # Get row count from in-memory structure
            row_count = 0
            for t in structure.get("tables", []):
                if t.get("id") == f"t_{t_idx}":
                    row_count = len(t.get("rows", []))
                    break

            # Add each row variable to data_dict
            for r_idx in range(row_count):
                var_name = f"{selection.variable_name}||{r_idx + 1}"
                data_dict[var_name] = context_val
        else:
            data_dict[selection.variable_name] = context_val

    # Apply formatting
    # First, handle columns
    for selection in col_assignments:
        parts = selection.id.split("_")
        t_idx = int(parts[1])
        col_idx = int(parts[3])
        if t_idx < len(doc.tables):
            table = doc.tables[t_idx]
            for r_idx, row in enumerate(table.rows):
                if col_idx < len(row.cells):
                    cell_id = f"t_{t_idx}_r_{r_idx}_c_{col_idx}"
                    # Only apply column variable if this specific cell hasn't been overridden
                    if cell_id not in cell_overrides:
                        row.cells[col_idx].text = f"{{{{ {selection.variable_name}_{r_idx} }}}}"
                        # Need to add this specific variable to data_dict to be complete? The user just wanted the column's var.
                        # We'll stick to the base column variable name mapping to context to keep data.json flat and clean.

    # Then handle individual elements
    for selection in request.selections:
        if selection.id.startswith("p_"):
            p_idx = int(selection.id.split("_")[1])
            if p_idx < len(doc.paragraphs):
                doc.paragraphs[p_idx].text = f"{{{{ {selection.variable_name} }}}}"
        elif selection.id.startswith("t_") and "col" not in selection.id:
            parts = selection.id.split("_")
            t_idx, r_idx, c_idx = int(parts[1]), int(parts[3]), int(parts[5])
            if t_idx < len(doc.tables):
                table = doc.tables[t_idx]
                if r_idx < len(table.rows) and c_idx < len(table.rows[r_idx].cells):
                    table.rows[r_idx].cells[c_idx].text = f"{{{{ {selection.variable_name} }}}}"

    output_path = os.path.join(UPLOAD_DIR, f"{request.doc_id}_templated.docx")
    doc.save(output_path)

    # Save templated document to data folder as doc_template.docx
    data_dir = os.path.join(os.path.dirname(__file__), "..", "data")
    os.makedirs(data_dir, exist_ok=True)
    template_path = os.path.join(data_dir, "doc_template.docx")
    doc.save(template_path)

    # Save data.json in the web folder
    data_json_path = os.path.join(
        os.path.dirname(__file__), "..", "web", "data.json")
    with open(data_json_path, "w") as f:
        json.dump(data_dict, f, indent=4)

    return {"message": "Document processed and saved to data/doc_template.docx", "mapping": data_dict}


@app.get("/download/{doc_id}")
async def download_document(doc_id: str):
    output_path = os.path.join(UPLOAD_DIR, f"{doc_id}_templated.docx")
    if not os.path.exists(output_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(output_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="templated_document.docx")


@app.get("/")
async def read_index():
    return FileResponse("web/index.html")

if os.path.exists("web"):
    app.mount("/static", StaticFiles(directory="web"), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
